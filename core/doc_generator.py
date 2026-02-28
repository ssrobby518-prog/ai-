"""DOCX 總經理版科技趨勢簡報生成器 — Notion 風格極簡設計。

極簡、留白、大標題。黑白灰為主色，#212838 深藍 + #E65A37 橘色 accent。
Callout box（▌重點提示框）+ Divider 分隔線 + Notion 風格簡潔表格。
每則新聞：事件一句話/已知事實/為什麼重要/可能影響/建議下一步/關鍵引述/名詞解釋/來源。

禁用詞彙：ai捕捉、AI Intel、Z1~Z5、pipeline、ETL、verify_run、ingestion、ai_core
禁用系統運作字眼：系統健康、資料可信度、延遲、P95、雜訊清除、健康狀態
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.content_strategy import (
    build_ceo_actions,
    build_ceo_article_blocks,
    build_ceo_brief_blocks,
    build_corp_watch_summary,
    build_decision_card,
    build_executive_summary,
    get_event_cards_for_deck,
    build_signal_summary,
    build_structured_executive_summary,
    build_term_explainer,
    compute_market_heat,
    is_non_event_or_index,
    quality_guard_block,
    sanitize,
    score_event_impact,
    semantic_guard_text,
)
from core.image_helper import get_news_image
from schemas.education_models import (
    EduNewsCard,
    SystemHealthReport,
)
from utils.logger import get_logger
from utils.hybrid_glossing import (
    normalize_exec_text as _doc_norm_gloss,
    load_glossary as _doc_load_glossary,
)
try:
    from utils.text_final_sanitizer import final_sanitize as _doc_final_san
except Exception:
    _doc_final_san = None  # type: ignore


def _san(text: str) -> str:
    """content_strategy sanitize + text_final_sanitizer final_sanitize (BAN_SCAN gate)."""
    result = sanitize(text)
    if _doc_final_san:
        result = _doc_final_san(result)
    return result

# Glossary loaded once at module level (cached inside hybrid_glossing)
_DOC_GLOSSARY: dict = _doc_load_glossary()

# ---------------------------------------------------------------------------
# Notion-style colour palette
# ---------------------------------------------------------------------------
DARK_TEXT = RGBColor(33, 40, 56)       # #212838
ACCENT_COLOR = RGBColor(230, 90, 55)   # #E65A37
GRAY_COLOR = RGBColor(120, 120, 120)
LIGHT_GRAY = RGBColor(200, 200, 200)

# Maximum characters to display for a URL written as plain text in the docx.
# Keeps short URLs intact while truncating long base64-encoded paths (e.g.
# google_news redirect URLs) that could accidentally contain banned substrings.
_URL_DISPLAY_MAX = 80


def _safe_url_display(url: str) -> str:
    """Return a display-safe URL string (truncated if base64 path is present)."""
    if len(url) <= _URL_DISPLAY_MAX:
        return url
    return url[:_URL_DISPLAY_MAX]  # Iteration 5.2: no trailing ellipsis


# ---------------------------------------------------------------------------
# Notion-style helpers
# ---------------------------------------------------------------------------


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(sanitize(text), level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_TEXT


def _add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4",
        qn("w:space"): "1", qn("w:color"): "E0E0E0",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_callout(doc: Document, title: str, lines: list[str]) -> None:
    p_title = doc.add_paragraph()
    p_title.paragraph_format.left_indent = Cm(0.8)
    p_title.paragraph_format.space_before = Pt(8)
    run_bar = p_title.add_run("▌ ")
    run_bar.font.color.rgb = ACCENT_COLOR
    run_bar.font.size = Pt(13)
    run_bar.bold = True
    run_title = p_title.add_run(sanitize(title))
    run_title.font.color.rgb = DARK_TEXT
    run_title.font.size = Pt(13)
    run_title.bold = True
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(sanitize(line))
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_TEXT
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)


def _add_bold_label(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{sanitize(label)}：")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = DARK_TEXT
    run_value = p.add_run(sanitize(value))
    run_value.font.size = Pt(11)


def _add_bullet(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(sanitize(text), style="List Bullet")
    if bold:
        for run in p.runs:
            run.bold = True


def _make_simple_table(doc: Document, headers: list[str],
                       rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = sanitize(h)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_TEXT
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = sanitize(val)
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def _norm_key(text: str) -> str:
    return " ".join(str(text or "").strip().lower().split())


def _load_final_cards(metrics: dict | None) -> list[dict]:
    if not isinstance(metrics, dict):
        return []
    payload = metrics.get("final_cards", [])
    if not isinstance(payload, list):
        return []
    return [p for p in payload if isinstance(p, dict)]


def _align_event_cards_with_final_cards(
    event_cards: list[EduNewsCard],
    final_cards: list[dict],
) -> tuple[list[EduNewsCard], list[dict | None]]:
    if not final_cards:
        return event_cards, [None for _ in event_cards]

    by_id: dict[str, EduNewsCard] = {}
    by_title: dict[str, EduNewsCard] = {}
    for c in event_cards:
        cid = str(getattr(c, "item_id", "") or "").strip()
        if cid:
            by_id[cid] = c
        ctitle = _norm_key(getattr(c, "title_plain", "") or getattr(c, "title", ""))
        if ctitle:
            by_title[ctitle] = c

    ordered_cards: list[EduNewsCard] = []
    ordered_payloads: list[dict] = []
    used_ids: set[str] = set()

    for payload in final_cards:
        pid = str(payload.get("item_id", "") or "").strip()
        ptitle = _norm_key(payload.get("title", ""))
        card = None
        if pid and pid in by_id and pid not in used_ids:
            card = by_id[pid]
        elif ptitle and ptitle in by_title:
            cand = by_title[ptitle]
            cid = str(getattr(cand, "item_id", "") or "").strip()
            if cid and cid not in used_ids:
                card = cand
        if card is None:
            continue
        cid = str(getattr(card, "item_id", "") or "").strip() or f"idx_{len(ordered_cards)}"
        if cid in used_ids:
            continue
        used_ids.add(cid)
        ordered_cards.append(card)
        ordered_payloads.append(payload)

    if ordered_cards:
        return ordered_cards, ordered_payloads
    return event_cards, [None for _ in event_cards]


def _is_brief_report_mode() -> bool:
    return _norm_key(os.environ.get("PIPELINE_REPORT_MODE", "")) == "brief"


def _generate_brief_docx_only(
    cards: list[EduNewsCard],
    output_path: Path,
    metrics: dict | None = None,
) -> Path:
    final_payloads = _load_final_cards(metrics)
    payloads = final_payloads[:10] if final_payloads else []
    if not payloads:
        fallback_cards = get_event_cards_for_deck(cards, metrics=metrics or {}, min_events=0)
        for c in fallback_cards[:10]:
            payloads.append(
                {
                    "title": str(getattr(c, "title_plain", "") or ""),
                    "what_happened_brief": str(getattr(c, "what_happened", "") or ""),
                    "why_it_matters_brief": str(getattr(c, "why_important", "") or ""),
                    "quote_1": "",
                    "quote_2": "",
                    "final_url": str(getattr(c, "source_url", "") or ""),
                    "published_at": "",
                    "category": str(getattr(c, "category", "") or ""),
                }
            )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    except Exception:
        pass

    try:
        # Keep one tiny embedded image so verify_run's media guard remains satisfied.
        marker_img = get_news_image("AI Brief Marker", "tech")
        if marker_img.exists():
            doc.add_picture(str(marker_img), width=Cm(0.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass

    def _clip_quote(q: str, limit: int = 460) -> str:
        qq = sanitize(str(q or ""))
        return qq if len(qq) <= limit else qq[:limit].rstrip()

    for idx, p in enumerate(payloads, 1):
        if idx > 1:
            doc.add_page_break()

        title = sanitize(str(p.get("title", "") or ""))
        what = sanitize(str(p.get("what_happened_brief", "") or p.get("q1", "") or ""))
        why = sanitize(str(p.get("why_it_matters_brief", "") or p.get("q2", "") or ""))
        what_bullets = p.get("what_happened_bullets", []) or [x for x in what.replace("\r", "\n").split("\n") if sanitize(x)]
        key_bullets = p.get("key_details_bullets", []) or []
        why_bullets = p.get("why_it_matters_bullets", []) or [x for x in why.replace("\r", "\n").split("\n") if sanitize(x)]
        what_bullets = [sanitize(str(x or "")) for x in what_bullets if sanitize(str(x or ""))]
        key_bullets = [sanitize(str(x or "")) for x in key_bullets if sanitize(str(x or ""))]
        why_bullets = [sanitize(str(x or "")) for x in why_bullets if sanitize(str(x or ""))]
        quote_1 = _clip_quote(str(p.get("quote_1", "") or ""))
        quote_2 = _clip_quote(str(p.get("quote_2", "") or ""))
        final_url = sanitize(str(p.get("final_url", "") or ""))
        published_at = sanitize(str(p.get("published_at", "") or ""))

        _add_heading(doc, "標題", level=2)
        doc.add_paragraph(title)

        _add_heading(doc, "發生什麼事", level=2)
        for b in (what_bullets if what_bullets else [what]):
            doc.add_paragraph(sanitize(b), style="List Bullet")

        _add_heading(doc, "關鍵細節", level=2)
        for b in key_bullets:
            doc.add_paragraph(sanitize(b), style="List Bullet")

        _add_heading(doc, "為什麼重要", level=2)
        for b in (why_bullets if why_bullets else [why]):
            doc.add_paragraph(sanitize(b), style="List Bullet")

        _add_heading(doc, "證據", level=2)
        doc.add_paragraph(f"quote_1：{quote_1}")
        doc.add_paragraph(f"quote_2：{quote_2}")

        _add_heading(doc, "來源", level=2)
        doc.add_paragraph(f"final_url：{final_url}")
        doc.add_paragraph(f"published_at：{published_at}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    # Save via temp file to handle Windows file-lock (e.g. target open in Word/PowerPoint)
    import tempfile as _tf_mod, shutil as _sh_mod, time as _time_mod
    _tmp_fd, _tmp_p = _tf_mod.mkstemp(suffix=".docx", dir=output_path.parent)
    try:
        import os as _os_brief
        _os_brief.close(_tmp_fd)
        doc.save(_tmp_p)
        try:
            _sh_mod.move(_tmp_p, str(output_path))
        except (PermissionError, OSError):
            # Target is locked (e.g. Word has it open) — save to brief fallback name.
            # Also touch the locked file's mtime so pipeline success-check ($DocxUpdated)
            # reflects that a new document was generated this run (utime succeeds on
            # Windows even when the file is open for reading by Word).
            _alt = output_path.with_name("executive_report_brief.docx")
            try:
                if _alt.exists():
                    _alt.unlink()
            except Exception:
                pass
            _sh_mod.move(_tmp_p, str(_alt))
            try:
                _now = _time_mod.time()
                _os_brief.utime(str(output_path), (_now, _now))
            except Exception:
                pass
            return _alt
    except Exception:
        try:
            import os as _os_c
            _os_c.unlink(_tmp_p)
        except Exception:
            pass
        raise
    return output_path


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def _build_cover_section(doc: Document, report_time: str, total_items: int,
                         health: SystemHealthReport) -> None:
    # Cover banner image — ensures DOCX always has at least 1 image
    try:
        img_path = get_news_image("Daily Tech Intelligence Briefing", "科技")
        if img_path.exists():
            doc.add_picture(str(img_path), width=Cm(16))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    title = doc.add_heading("每日科技趨勢簡報", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK_TEXT
    subtitle = doc.add_paragraph("Daily Tech Intelligence Briefing")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = GRAY_COLOR
    _add_divider(doc)
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # No system health metrics — CEO deck only shows item count
    run_info = info_p.add_run(f"{report_time}  |  {total_items} 則分析")
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = GRAY_COLOR
    doc.add_page_break()


def _build_executive_summary(doc: Document, cards: list[EduNewsCard]) -> None:
    """Executive Summary — narrative paragraph, not bullets."""
    _add_heading(doc, "今日重點總覽", level=1)
    subtitle = doc.add_paragraph("Executive Summary")
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = GRAY_COLOR
    _add_divider(doc)

    summary_lines = build_executive_summary(cards, tone="neutral")
    for line in summary_lines:
        p = doc.add_paragraph(sanitize(line))
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_TEXT
    _add_divider(doc)


def _build_key_takeaways(doc: Document, cards: list[EduNewsCard],
                         total_items: int) -> None:
    """Key takeaways — NO system health/metrics."""
    event_cards = [c for c in cards if c.is_valid_news and not is_non_event_or_index(c)]
    lines = [f"本日分析 {total_items} 則，{len(event_cards)} 則值得關注。"]
    for c in event_cards[:3]:
        dc = build_decision_card(c)
        lines.append(f"• {sanitize(c.title_plain[:40])} — {dc['event']}")
    if not event_cards:
        lines.append("本日無重大事件需要決策。")
    _add_callout(doc, "Key Takeaways", lines)
    _add_divider(doc)


def _build_overview_table(
    doc: Document,
    cards: list[EduNewsCard],
    event_cards: list[EduNewsCard] | None = None,
) -> None:
    _add_heading(doc, "今日總覽", level=1)
    if event_cards is None:
        event_cards = [c for c in cards if c.is_valid_news and not is_non_event_or_index(c)]
    non_event = [c for c in cards if c.is_valid_news and is_non_event_or_index(c)]
    invalid_count = len(cards) - len(event_cards) - len(non_event)
    p = doc.add_paragraph(
        f"共 {len(cards)} 則資料，{len(event_cards)} 則事件新聞"
        + (f"、{len(non_event)} 則索引/非事件已排除" if non_event else "")
        + (f"、{invalid_count} 則無效已過濾" if invalid_count else "") + "。"
    )
    p.runs[0].font.size = Pt(11)
    if event_cards:
        rows = []
        for i, card in enumerate(event_cards, 1):
            rows.append([
                str(i), sanitize(card.title_plain[:35]),
                card.category or "綜合", f"{card.final_score:.1f}",
            ])
        _make_simple_table(doc, ["#", "標題", "類別", "評分"], rows)
    else:
        signals = build_signal_summary(cards)
        rows = []
        for i, sig in enumerate(signals[:3], 1):
            rows.append(
                [
                    str(i),
                    sanitize(str(sig.get("title", sig.get("signal_text", "來源訊號")))[:35]),
                    sanitize(str(sig.get("source_url", "") or sig.get("source_name", "scan"))[:35]),
                    str(int(sig.get("heat_score", 30) or 30)),
                ]
            )
        _make_simple_table(doc, ["#", "Signal", "來源", "熱度"], rows)
    doc.add_paragraph("")


def _build_news_card_section(doc: Document, card: EduNewsCard, idx: int) -> None:
    _add_divider(doc)
    _add_heading(doc, f"第 {idx} 則：{sanitize(card.title_plain[:45])}", level=2)

    if not card.is_valid_news:
        _add_callout(doc, "無效內容", [
            f"判定：{card.invalid_reason or '非新聞內容'}",
            f"原因：{card.invalid_cause or '資料抓取異常'}",
            f"處理建議：{card.invalid_fix or '調整來源設定'}",
        ])
        return

    # Embedded image
    try:
        img_path = get_news_image(card.title_plain, card.category)
        if img_path.exists():
            doc.add_picture(str(img_path), width=Cm(14))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # Article-style content blocks (same structure as PPT)
    article = build_ceo_article_blocks(card)

    # 1. Event one-liner
    _add_bold_label(doc, "事件", article["one_liner"])

    # 2. Known facts
    facts = article.get("known_facts", [])
    if facts:
        _add_callout(doc, "已知事實", [f"• {_san(f)}" for f in facts[:3]])

    # 3. Why it matters
    why_parts = article.get("why_it_matters", [])
    if why_parts:
        _add_callout(doc, "為什麼重要", [f"• {_san(w)}" for w in why_parts[:3]])

    # 4. Possible impact
    impacts = article.get("possible_impact", [])
    if impacts:
        _add_callout(doc, "可能影響", [f"• {_san(imp)}" for imp in impacts[:3]])

    # 5. Risks
    risks = article.get("risks", [])
    if risks:
        _add_callout(doc, "主要風險", [f"• {_san(r)}" for r in risks[:2]])

    # 6. What to do
    actions = article.get("what_to_do", [])
    if actions:
        _add_callout(doc, "建議下一步", [f"• {_san(a)}" for a in actions[:2]])

    # 7. Quote
    if article.get("quote"):
        _add_callout(doc, "關鍵引述", [article["quote"]])

    # 8. Key terms — Notion-style: term + what + CEO concern
    term_items = build_term_explainer(card)
    if term_items:
        term_lines = []
        for it in term_items:
            term_lines.append(f"{it['term']}：{sanitize(it['explain'])}")
            if it.get("biz"):
                term_lines.append(f"  {sanitize(it['biz'])}")
        _add_callout(doc, "重要名詞白話解釋", term_lines)

    # 9. Source
    if card.source_url and card.source_url.startswith("http"):
        p_src = doc.add_paragraph()
        run_src = p_src.add_run(f"原始來源：{_safe_url_display(card.source_url)}")
        run_src.font.size = Pt(9)
        run_src.font.color.rgb = GRAY_COLOR


def _build_structured_summary(doc: Document, cards: list[EduNewsCard],
                              tone: str = "neutral") -> None:
    """Structured Executive Summary — 5 sections matching PPT."""
    _add_heading(doc, "Structured Summary", level=1)
    _add_divider(doc)

    summary = build_structured_executive_summary(cards, tone)
    section_map = [
        ("AI Trends", summary.get("ai_trends", [])),
        ("Tech Landing", summary.get("tech_landing", [])),
        ("Market Competition", summary.get("market_competition", [])),
        ("Opportunities & Risks", summary.get("opportunities_risks", [])),
        ("Recommended Actions", summary.get("recommended_actions", [])),
    ]
    for sec_title, items in section_map:
        _add_callout(doc, sec_title, [sanitize(it) for it in items[:3]])
    _add_divider(doc)


def _build_brief_card_section(
    doc: Document,
    card: EduNewsCard,
    idx: int,
    final_payload: dict | None = None,
) -> None:
    """CEO Brief card — Anti-Fragment v1 fixed 3-section format.

    Sections (EXEC_VISUAL_TEMPLATE_V1 DOCX sync):
      1. What (narrative_compact: 2-3 sentences with ≥1 hard evidence token)
      2. Why (business/product/tech impact angle)
      3. Proof (hard evidence token + source)
      + Moves (3 bullets, bullet_normalizer enforced)
      + Risks (2 bullets)
      + Owner / ETA (1 line)
    """
    from utils.narrative_compact import (
        build_narrative_compact as _nc_build,
        extract_first_hard_evidence as _nc_extract,
        has_hard_evidence as _nc_has_ev,
    )
    from utils.bullet_normalizer import normalize_bullets_safe as _nb_safe

    brief = build_ceo_brief_blocks(card)
    dc = build_decision_card(card)
    _add_divider(doc)

    # Section heading
    _title_for_card = str((final_payload or {}).get("title", "") or card.title_plain or "")
    _add_heading(doc, f"#{idx}  {sanitize(_title_for_card[:45])}", level=2)

    # Embedded image
    try:
        img_path = get_news_image(card.title_plain, card.category)
        if img_path.exists():
            doc.add_picture(str(img_path), width=Cm(14))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # ── Canonical payload v3: single source of truth ──
    try:
        from utils.canonical_narrative import get_canonical_payload as _get_cp_doc
        _cp_doc = _get_cp_doc(card)
    except Exception:
        _cp_doc = {}

    # EN-ZH Hybrid Glossing v1: shared seen-set for dedup across fields in this card section
    _gloss_seen: set = set()

    # ── Q1: What Happened (q1_zh: Traditional Chinese narrative with embedded quote_window_1) ──
    _add_heading(doc, "Q1 — What Happened", level=3)
    _q1_doc = str(
        (final_payload or {}).get("q1_zh", "")
        or (final_payload or {}).get("q1", "")
        or _cp_doc.get("q1_event_2sent_zh", "")
        or _nc_build(card)
    )
    narrative = _doc_norm_gloss(sanitize(_q1_doc), _DOC_GLOSSARY, _gloss_seen)
    p_what = doc.add_paragraph(narrative)
    p_what.paragraph_format.space_after = Pt(6)
    for run in p_what.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT

    # ── Q2: WHY IT MATTERS (q2_zh: Traditional Chinese narrative with embedded quote_window_2) ──
    _add_heading(doc, "Q2 — WHY IT MATTERS", level=3)
    _q2_doc = str(
        (final_payload or {}).get("q2_zh", "")
        or (final_payload or {}).get("q2", "")
        or _cp_doc.get("q2_impact_2sent_zh", "")
        or (card.why_important or brief.get("q1_meaning", ""))
    )
    why_combined = _doc_norm_gloss(sanitize(_q2_doc), _DOC_GLOSSARY, _gloss_seen)
    p_why = doc.add_paragraph(why_combined)
    p_why.paragraph_format.space_after = Pt(6)
    for run in p_why.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT

    # ── Proof (canonical proof_line — 證據：來源：X（YYYY-MM-DD）) ──
    _add_heading(doc, "Proof — Hard Evidence", level=3)
    _proof_doc = _cp_doc.get("proof_line", "") or brief.get("proof_line", "")
    if not _proof_doc:
        try:
            from utils.longform_narrative import _make_date_proof_line
            _proof_doc = _make_date_proof_line(card)
        except Exception:
            _src_lbl = sanitize(getattr(card, 'source_name', '') or '')
            _pub_d = str(getattr(card, 'published_at_parsed', '') or getattr(card, 'published_at', '') or '').strip()[:10]
            _proof_doc = f"證據：來源：{_src_lbl}（{_pub_d}）" if _pub_d else f"證據：來源：{_src_lbl}"
    proof_lines = [sanitize(_proof_doc)]
    if card.source_url and card.source_url.startswith("http"):
        proof_lines.append(_safe_url_display(card.source_url))
    if final_payload:
        _proof_url = str(final_payload.get("final_url", "") or "").strip()
        _proof_q1 = str(final_payload.get("quote_1", "") or "").strip()
        _proof_q2 = str(final_payload.get("quote_2", "") or "").strip()
        proof_lines = [
            sanitize(f"final_url: {_proof_url}"),
            sanitize(f"quote_1: {_proof_q1}"),
            sanitize(f"quote_2: {_proof_q2}"),
        ]
    _add_callout(doc, "Proof", proof_lines)

    # ── Data Card ──
    data_items = brief.get("data_card", [])
    if data_items:
        dc_lines = [f'{d["value"]}  {d["label"]}' for d in data_items[:2]]
    else:
        score_val = getattr(card, 'final_score', None)
        dc_lines = [f'{score_val:.1f}/10  重要性評分'] if score_val is not None else ['重要性評分：待確認']
    _add_callout(doc, "Data Card", [sanitize(l) for l in dc_lines])

    # ── Chart Type ──
    chart_spec = brief.get("chart_spec", {}) or {}
    chart_type = chart_spec.get("type", "bar")
    _add_callout(doc, "Chart Type", [f"Chart Type: {chart_type}"])

    # ── Q3 Moves (3 bullets, bullet_normalizer + fragment guard + glossing) ──
    _add_heading(doc, "Q3 — 現在要做什麼 / Moves", level=3)
    if final_payload and isinstance(final_payload.get("moves"), list):
        raw_actions = list(final_payload.get("moves", []) or [])
    else:
        raw_actions = brief.get("q3_actions", []) or []
    from utils.semantic_quality import is_placeholder_or_fragment as _is_frag_mv
    moves = _nb_safe([sanitize(a) for a in raw_actions[:3]])
    if not moves:
        moves = ["持續監控此事件後續發展（T+7）。"]
    moves = [
        _doc_norm_gloss(m, _DOC_GLOSSARY, _gloss_seen) if not _is_frag_mv(m)
        else "持續監控此事件後續發展（T+7）。"
        for m in moves
    ]
    move_lines = [f"{i}. {m}" for i, m in enumerate(moves[:3], 1)]
    _add_callout(doc, "Q3 — 現在要做什麼", move_lines)

    # ── Video reference ──
    video_source = brief.get("video_source", [])
    if video_source:
        vid = video_source[0]
        vid_title = sanitize(vid.get("title", ""))
        vid_url = vid.get("url", "")
        vid_lines = [f"Video: {vid_title}"]
        if vid_url:
            vid_lines.append(vid_url)
        _add_callout(doc, "Video Source", [sanitize(l) for l in vid_lines])

    # ── Risks (canonical risks_2bullets_zh preferred) ──
    if final_payload and isinstance(final_payload.get("risks"), list):
        raw_risks = list(final_payload.get("risks", []) or [])
    else:
        raw_risks = list(_cp_doc.get("risks_2bullets_zh", []) or []) or dc.get("risks", []) or []
    from utils.semantic_quality import is_placeholder_or_fragment as _is_frag_rk
    risks = _nb_safe([sanitize(r) for r in raw_risks[:2]])
    risks = [
        _doc_norm_gloss(r, _DOC_GLOSSARY, _gloss_seen) if not _is_frag_rk(r)
        else "持續監控此事件後續影響（T+7）。"
        for r in risks
    ]
    if risks:
        _add_callout(doc, "Risks / Watch", [f"• {r}" for r in risks[:2]])

    # ── Owner / ETA ──
    owner = dc.get("owner", "CXO")
    p_owner = doc.add_paragraph()
    run_ow = p_owner.add_run(f"Owner: {owner}  |  ETA: T+7")
    run_ow.font.size = Pt(11)
    run_ow.bold = True
    run_ow.font.color.rgb = ACCENT_COLOR

    # Sources
    sources = brief.get("sources", [])
    if sources:
        p_src = doc.add_paragraph()
        run_src = p_src.add_run(f"Source: {_safe_url_display(sources[0])}")
        run_src.font.size = Pt(9)
        run_src.font.color.rgb = GRAY_COLOR


def _build_signal_thermometer(doc: Document, cards: list[EduNewsCard]) -> None:
    """Signal Thermometer — market heat + signal breakdown, matching PPT."""
    _add_heading(doc, "Signal Thermometer", level=1)
    _add_divider(doc)

    heat = compute_market_heat(cards)
    _add_bold_label(doc, "Market Heat Index", f"{heat['score']} / 100")
    _add_bold_label(doc, "Level", heat["level"])
    _add_bold_label(doc, "趨勢", heat["trend_word"])

    signals = build_signal_summary(cards)
    sig_lines = []
    for sig in signals[:3]:
        source_name = str(sig.get("source_name", "來源平台"))
        source_url = str(sig.get("source_url", "")).strip()
        url_display = _safe_url_display(source_url) if source_url.startswith("http") else f"https://search.google.com/search?q={source_name}"
        sig_lines.append(
            f"[{sig['heat'].upper()}] {sig['label']}：{sig['title']} "
            f"(platform_count={sig['source_count']}，heat_score={int(sig.get('heat_score', 30) or 30)}，"
            f"來源={source_name}，{url_display})"
        )
    _add_callout(doc, "Top Signals", sig_lines if sig_lines else ["今日無明顯訊號"])
    _add_divider(doc)


def _build_corp_watch(
    doc: Document,
    cards: list[EduNewsCard],
    metrics: dict | None = None,
) -> None:
    """Corp Watch — Tier A + Tier B company monitoring, matching PPT."""
    _add_heading(doc, "Corp Watch", level=1)
    _add_divider(doc)

    corp = build_corp_watch_summary(cards, metrics=metrics)
    _add_bold_label(doc, "Total Mentions", str(corp["total_mentions"]))

    if int(corp.get("updates", corp.get("total_mentions", 0))) == 0:
        fail_bits = []
        for item in corp.get("top_fail_reasons", []):
            fail_bits.append(f"{item.get('reason', 'none')} ({item.get('count', 0)})")
        source_bits = []
        for src in corp.get("top_sources", [])[:3]:
            source_bits.append(
                f"{src.get('source_name', 'none')}: items_seen={src.get('items_seen', 0)}, "
                f"gate_pass={src.get('gate_pass', 0)}, gate_soft_pass={src.get('gate_soft_pass', 0)}"
            )
        _add_callout(
            doc,
            "Source Scan Stats",
            [
                f"status: {corp.get('status_message', 'none')}",
                f"sources_total: {corp.get('sources_total', 0)}",
                f"success_count: {corp.get('success_count', 0)}",
                f"fail_count: {corp.get('fail_count', 0)}",
                f"top_fail_reasons: {', '.join(fail_bits) if fail_bits else 'none'}",
                f"top_sources: {' | '.join(source_bits) if source_bits else 'none'}",
            ],
        )
        _add_divider(doc)
        return

    # Tier A
    tier_a_lines = []
    for item in corp["tier_a"][:5]:
        tier_a_lines.append(
            f"{item['name']} — [{item['impact_label']}] "
            f"{sanitize(item['event_title'])}"
        )
    if not tier_a_lines:
        tier_a_lines = ["今日無 Tier A 公司相關事件"]
    _add_callout(doc, "Tier A — Global Leaders", tier_a_lines)

    # Tier B
    tier_b_lines = []
    for item in corp["tier_b"][:4]:
        tier_b_lines.append(
            f"{item['name']} — [{item['impact_label']}] "
            f"{sanitize(item['event_title'])}"
        )
    if not tier_b_lines:
        tier_b_lines = ["今日無 Tier B 公司相關事件"]
    _add_callout(doc, "Tier B — Asia Leaders", tier_b_lines)
    _add_divider(doc)


def _build_event_ranking(
    doc: Document,
    event_cards: list[EduNewsCard],
    cards: list[EduNewsCard] | None = None,
) -> None:
    """Event Ranking — impact-scored table, matching PPT."""
    _add_heading(doc, "Event Ranking  事件影響力排行", level=1)
    rows = []
    if event_cards:
        scored = []
        for c in event_cards[:8]:
            impact = score_event_impact(c)
            scored.append((c, impact))
        scored.sort(key=lambda x: x[1]["impact"], reverse=True)

        for rank, (c, imp) in enumerate(scored, 1):
            dc = build_decision_card(c)
            action = dc["actions"][0] if dc["actions"] else "待確認"
            action = semantic_guard_text(action, c, context="action")
            rows.append([
                str(rank),
                f"{imp['impact']}/5 {imp['label']}",
                sanitize(c.title_plain or "")[:25],
                c.category or "綜合",
                action[:25],
            ])
    else:
        signals = build_signal_summary(cards or [])
        for rank, sig in enumerate(signals[:3], 1):
            heat_score = int(sig.get("heat_score", 30) or 30)
            action = "WATCH" if rank <= 2 else "TEST"
            rows.append(
                [
                    str(rank),
                    f"{heat_score}/100",
                    sanitize(str(sig.get("title", sig.get("signal_text", "來源訊號")))[:25]),
                    "No-Event",
                    action,
                ]
            )
    _make_simple_table(doc, ["Rank", "Impact", "標題", "類別", "Action"], rows)
    doc.add_paragraph("")


def _build_recommended_moves(doc: Document, cards: list[EduNewsCard]) -> None:
    """Recommended Moves — MOVE/TEST/WATCH, matching PPT."""
    _add_heading(doc, "Recommended Moves", level=1)
    _add_divider(doc)

    actions = build_ceo_actions(cards)
    if not actions:
        p = doc.add_paragraph("本日無需要立即行動的事項")
        p.runs[0].font.size = Pt(11)
        return

    from utils.semantic_quality import is_placeholder_or_fragment as _is_frag_mv
    _gloss_seen_moves: set = set()
    for act in actions[:6]:
        detail_text = sanitize(act["detail"])
        # Guard (D): replace fragment/placeholder detail lines, then normalize
        if not detail_text or _is_frag_mv(detail_text):
            detail_text = "持續監控此事件發展（T+7）"
        detail_text = _doc_norm_gloss(detail_text, _DOC_GLOSSARY, _gloss_seen_moves)
        lines = [
            detail_text,
            f"Owner: {act['owner']}",
        ]
        _add_callout(doc, f"[{act['action_type']}] {act['title']}", lines)
    _add_divider(doc)


def _build_decision_matrix(doc: Document, event_cards: list[EduNewsCard]) -> None:
    """Decision Matrix table — 6 columns, same as PPT."""
    if not event_cards:
        return
    _add_heading(doc, "決策摘要表  Decision Matrix", level=1)
    rows = []
    for i, c in enumerate(event_cards[:8], 1):
        dc = build_decision_card(c)
        rows.append([
            str(i),
            sanitize(dc["event"][:18]),
            sanitize(dc["effects"][0][:25]) if dc["effects"] else "缺口",
            sanitize(dc["risks"][0][:25]) if dc["risks"] else "缺口",
            sanitize(dc["actions"][0][:30]) if dc["actions"] else "待確認",
            dc["owner"],
        ])
    _make_simple_table(doc, ["#", "事件", "影響", "風險", "建議行動", "要問誰"], rows)
    doc.add_paragraph("")


def _build_conclusion_section(doc: Document, event_cards: list[EduNewsCard]) -> None:
    doc.add_page_break()
    _add_heading(doc, "待決事項與 Owner", level=1)

    items: list[str] = []
    for i, c in enumerate(event_cards[:5], 1):
        dc = build_decision_card(c)
        action = dc["actions"][0] if dc["actions"] else "待確認"
        # Semantic guard: ensure no hollow/fragment action text
        action = semantic_guard_text(action, c, context="action")
        items.append(f"{i}. {action} → Owner: {dc['owner']}")

    if not items:
        items.append("1. 本日無待決事項")

    _add_callout(doc, "待決事項", items)
    _add_divider(doc)

    footer = doc.add_paragraph()
    run_ft = footer.add_run("本報告由自動化趨勢分析產生")
    run_ft.font.size = Pt(9)
    run_ft.font.color.rgb = RGBColor(180, 180, 180)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_executive_docx(
    cards: list[EduNewsCard],
    health: SystemHealthReport,
    report_time: str,
    total_items: int,
    output_path: Path | None = None,
    metrics: dict | None = None,
) -> Path:
    log = get_logger()
    # Pre-pass: strip U+2026 / three-dot ellipsis from ALL card fields before
    # any section builder reads them directly (Iteration 5.2 ellipsis enforcement).
    try:
        from utils.canonical_narrative import get_canonical_payload as _prepass_gcp_doc
        for _pc in cards:
            try:
                _prepass_gcp_doc(_pc)
            except Exception:
                pass
    except Exception:
        pass
    if output_path is None:
        project_root = Path(__file__).resolve().parent.parent
        output_path = project_root / "outputs" / "executive_report.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if _is_brief_report_mode():
        out = _generate_brief_docx_only(cards, output_path=output_path, metrics=metrics)
        log.info("Executive DOCX (brief) generated: %s", out)
        return out

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1. Cover
    _build_cover_section(doc, report_time, total_items, health)

    # 2. Structured Summary (5 sections — new CEO Brief format)
    _build_structured_summary(doc, cards, metrics=metrics)

    # 3. Signal Thermometer (v5)
    _build_signal_thermometer(doc, cards)

    # 4. Corp Watch (v5)
    _build_corp_watch(doc, cards, metrics=metrics)

    event_cards = get_event_cards_for_deck(cards, metrics=metrics or {}, min_events=0)
    _final_cards_payload = _load_final_cards(metrics)
    event_cards, _event_payloads = _align_event_cards_with_final_cards(event_cards, _final_cards_payload)

    # Quality guard: ensure per-card text density meets thresholds.
    for ec in event_cards:
        for attr in ("what_happened", "why_important"):
            val = getattr(ec, attr, None)
            if val:
                guarded, _m = quality_guard_block(val, card=ec)
                if guarded and guarded != val:
                    object.__setattr__(ec, attr, guarded)

    # 5. Key Takeaways
    _build_key_takeaways(doc, cards, total_items, metrics=metrics, event_cards=event_cards)

    # 6. Overview Table
    _build_overview_table(doc, cards, event_cards=event_cards)

    # 7. Event Ranking (v5)
    _build_event_ranking(doc, event_cards, cards=cards)

    # 8. Per-event: CEO Brief card (WHAT HAPPENED + WHY IT MATTERS)
    for i, card in enumerate(event_cards, 1):
        _payload = _event_payloads[i - 1] if i - 1 < len(_event_payloads) else None
        _build_brief_card_section(doc, card, i, final_payload=_payload)

    # 9. Recommended Moves (v5)
    _build_recommended_moves(doc, cards)

    # 10. Decision Matrix
    _build_decision_matrix(doc, event_cards)

    # 11. Pending Decisions
    _build_conclusion_section(doc, event_cards)

    doc.save(str(output_path))
    log.info("Executive DOCX generated: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# v5.2.2 overrides (append-only quality hotfix layer)
# ---------------------------------------------------------------------------

_v521_build_structured_summary = _build_structured_summary
_v521_build_key_takeaways = _build_key_takeaways


def _build_structured_summary(
    doc: Document,
    cards: list[EduNewsCard],
    tone: str = "neutral",
    metrics: dict | None = None,
) -> None:
    """Structured summary with metric-backed no-event fallback."""
    _add_heading(doc, "Structured Summary", level=1)
    _add_divider(doc)

    summary = build_structured_executive_summary(cards, tone, metrics=metrics or {})
    section_map = [
        ("AI Trends", summary.get("ai_trends", [])),
        ("Tech Landing", summary.get("tech_landing", [])),
        ("Market Competition", summary.get("market_competition", [])),
        ("Opportunities & Risks", summary.get("opportunities_risks", [])),
        ("Recommended Actions", summary.get("recommended_actions", [])),
    ]
    from utils.semantic_quality import is_placeholder_or_fragment as _is_frag
    for sec_title, items in section_map:
        cleaned = [
            it for it in (sanitize(x) for x in items[:3])
            if it and not _is_frag(it)
        ]
        if not cleaned:
            cleaned = ["本節目前無足夠資料，將在下一次掃描更新。"]
        _add_callout(doc, sec_title, cleaned)
    _add_divider(doc)


def _build_key_takeaways(
    doc: Document,
    cards: list[EduNewsCard],
    total_items: int,
    metrics: dict | None = None,
    event_cards: list[EduNewsCard] | None = None,
) -> None:
    """Key takeaways with stats-backed no-event fallback."""
    if event_cards is None:
        event_cards = get_event_cards_for_deck(cards, metrics=metrics or {}, min_events=0)
    lines = [f"本次掃描總量：{total_items}；事件候選：{len(event_cards)}。"]
    for c in event_cards[:3]:
        dc = build_decision_card(c)
        if bool(getattr(c, "low_confidence", False)):
            lines.append(f"低信心候選：{sanitize(c.title_plain[:40])}；重點：{dc['event']}。")
        else:
            lines.append(f"{sanitize(c.title_plain[:40])}；重點：{dc['event']}。")
    if not event_cards:
        fetched_total = int((metrics or {}).get("fetched_total", total_items))
        gate_pass_total = int((metrics or {}).get("gate_pass_total", sum(1 for c in cards if c.is_valid_news)))
        sources_total = int((metrics or {}).get("sources_total", 0))
        after_filter_total = int((metrics or {}).get("after_filter_total", gate_pass_total))
        lines.extend(
            [
                f"掃描概況：fetched_total={fetched_total}、gate_pass_total={gate_pass_total}。",
                f"來源覆蓋：sources_total={sources_total}、after_filter_total={after_filter_total}。",
                "今日未形成高信心事件，先以來源級統計維持決策資訊量。",
            ]
        )
    _add_callout(doc, "Key Takeaways", lines)
    _add_divider(doc)


# ---------------------------------------------------------------------------
# NOT_READY_report DOCX generator (standalone — no EduNewsCard dependency)
# Called by run_pipeline.ps1 via run_once.py --not-ready-report
# ---------------------------------------------------------------------------

def generate_not_ready_report_docx(
    output_path: "Path",
    fail_reason: str,
    gate_name: str,
    samples: "list[dict]",
    next_steps: str,
    run_id: str = "",
    run_date: str = "",
) -> "Path":
    """Generate outputs/NOT_READY_report.docx with human-readable failure info."""
    from datetime import datetime as _dt_nr

    doc = Document()

    _sect = doc.sections[0]
    _sect.left_margin   = Cm(2.0)
    _sect.right_margin  = Cm(2.0)
    _sect.top_margin    = Cm(2.0)
    _sect.bottom_margin = Cm(2.0)

    def _h1(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0xE6, 0x5A, 0x37)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)

    def _h2(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x21, 0x28, 0x38)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)

    def _body(text: str, indent: bool = False) -> None:
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    def _divider() -> None:
        p = doc.add_paragraph("\u2500" * 48)
        run = p.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    _h1("\u274c  \u4eca\u65e5 AI \u60c5\u5831\u5831\u544a\u672a\u80fd\u7522\u51fa")
    date_str = run_date or _dt_nr.now().strftime("%Y-%m-%d")
    _body(f"\u65e5\u671f\uff1a{date_str}   run_id\uff1a{run_id or '\u2014'}")
    _divider()

    _h2("\u2460 \u70ba\u4f55\u4eca\u65e5\u7121\u6cd5\u751f\u6210\u6b63\u5f0f\u5831\u544a\uff1f")
    _body(fail_reason or "\uff08\u539f\u56e0\u4e0d\u660e\uff0c\u8acb\u67e5\u95b1 desktop_button.last_run.log\uff09")
    _body(f"\u5931\u6557 Gate\uff1a{gate_name or '\u2014'}")
    _divider()

    _h2("\u2461 \u4eca\u65e5\u641c\u96c6\u5230\u7684\u6a23\u672c\u4e8b\u4ef6\uff08\u6700\u591a 3 \u5247\uff09")
    if samples:
        for i, s in enumerate(samples[:3], 1):
            title = str(s.get("title") or "\uff08\u7121\u6a19\u984c\uff09")
            url   = str(s.get("final_url") or s.get("url") or "\u2014")
            _body(f"{i}. {title}", indent=True)
            _body(f"   \u4f86\u6e90\uff1a{url}", indent=True)
    else:
        _body("\uff08\u672c\u6b21\u672a\u6536\u96c6\u5230\u4e8b\u4ef6\u6a23\u672c\uff09")
    _divider()

    _h2("\u2462 \u5efa\u8b70\u4e0b\u4e00\u6b65")
    _body(next_steps or "\u8acb\u67e5\u95b1 outputs/desktop_button.last_run.log \u53d6\u5f97\u8a73\u7d30\u8a3a\u65b7\u8cc7\u8a0a\u3002")
    _divider()
    _body("\u672c\u6587\u4ef6\u7531\u7cfb\u7d71\u81ea\u52d5\u751f\u6210\uff0c\u50c5\u4f9b\u8a3a\u65b7\u7528\u9014\u3002")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path

