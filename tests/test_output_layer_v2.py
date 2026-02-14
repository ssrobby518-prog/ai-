"""Output Layer v2 tests — PPTX + DOCX generation.

Covers:
- PPT generator produces valid .pptx with expected slide count
- DOCX generator produces valid .docx with expected sections
- Binary reports work with empty card list
- Integration: generate_binary_reports wrapper
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from schemas.education_models import EduNewsCard, SystemHealthReport
from schemas.models import (
    DeepAnalysisReport,
    ItemDeepDive,
    MergedResult,
    SchemaA,
    SchemaB,
    SchemaC,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_METRICS: dict = {
    "run_id": "test_run_v2",
    "total_items": 2,
    "enrich_success_rate": 80.0,
    "enrich_latency_p50": 4.0,
    "enrich_latency_p95": 11.0,
    "entity_noise_removed": 2,
    "total_runtime_seconds": 30.0,
    "enrich_fail_reasons": {"blocked": 1},
}


def _make_health() -> SystemHealthReport:
    return SystemHealthReport(
        success_rate=80.0,
        p50_latency=4.0,
        p95_latency=11.0,
        entity_noise_removed=2,
        total_runtime=30.0,
        run_id="test_run_v2",
        fail_reasons={"blocked": 1},
    )


def _make_valid_card(idx: int = 1) -> EduNewsCard:
    return EduNewsCard(
        item_id=f"news_{idx:03d}",
        is_valid_news=True,
        title_plain=f"Test News Title {idx}",
        what_happened="Something important happened in the tech world.",
        why_important="This affects the entire industry supply chain.",
        focus_action="Monitor subsequent announcements from major players.",
        metaphor="Like upgrading a factory assembly line — short-term pain, long-term gain.",
        fact_check_confirmed=["Fact A confirmed", "Fact B confirmed"],
        fact_check_unverified=["Claim X needs verification"],
        evidence_lines=["Evidence: original text excerpt here"],
        technical_interpretation="Technical analysis of the event spanning multiple dimensions.",
        derivable_effects=["Direct effect on market pricing"],
        speculative_effects=["Possible regulatory response"],
        observation_metrics=["Watch industry adoption rate"],
        action_items=["This week: research latest reports", "Next month: assess impact"],
        image_suggestions=["Image suggestion placeholder"],
        video_suggestions=["YouTube search: test topic"],
        reading_suggestions=["Google search: test analysis"],
        source_url="https://example.com/news",
        category="tech",
        signal_strength=0.8,
        final_score=8.0,
        source_name="TestSource",
    )


def _make_invalid_card() -> EduNewsCard:
    return EduNewsCard(
        item_id="invalid_001",
        is_valid_news=False,
        invalid_reason="System banner, not real news",
        title_plain="Invalid Content",
        what_happened="Scraper captured a login page.",
        why_important="Recognizing invalid content is important.",
        invalid_cause="blocked / anti-scraping",
        invalid_fix="Adjust scraping strategy",
        evidence_lines=["Original: You signed in with another tab"],
        source_url="N/A",
        signal_strength=0.0,
    )


def _make_cards() -> list[EduNewsCard]:
    return [_make_valid_card(1), _make_valid_card(2), _make_invalid_card()]


# ---------------------------------------------------------------------------
# PPT Generator Tests
# ---------------------------------------------------------------------------


class TestPptGenerator:
    def test_generates_pptx_file(self, tmp_path):
        from core.ppt_generator import generate_executive_ppt

        out = tmp_path / "test.pptx"
        result = generate_executive_ppt(
            cards=_make_cards(),
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=3,
            output_path=out,
        )
        assert result.exists()
        assert result.suffix == ".pptx"
        assert result.stat().st_size > 1000  # not empty

    def test_slide_count(self, tmp_path):
        from core.ppt_generator import generate_executive_ppt
        from pptx import Presentation

        out = tmp_path / "test.pptx"
        generate_executive_ppt(
            cards=_make_cards(),
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=3,
            output_path=out,
        )
        prs = Presentation(str(out))
        # Cover(1) + Contents(1) + Section1(1) + Summary(1)
        # + Section2(1) + 2 valid cards * 2 slides + 1 invalid card * 1 slide
        # + Section3(1) + Metrics(1) + Conclusion(1) = 13
        assert len(prs.slides) >= 10

    def test_empty_cards(self, tmp_path):
        from core.ppt_generator import generate_executive_ppt

        out = tmp_path / "empty.pptx"
        result = generate_executive_ppt(
            cards=[],
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=0,
            output_path=out,
        )
        assert result.exists()
        assert result.stat().st_size > 500


# ---------------------------------------------------------------------------
# DOCX Generator Tests
# ---------------------------------------------------------------------------


class TestDocxGenerator:
    def test_generates_docx_file(self, tmp_path):
        from core.doc_generator import generate_executive_docx

        out = tmp_path / "test.docx"
        result = generate_executive_docx(
            cards=_make_cards(),
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=3,
            output_path=out,
        )
        assert result.exists()
        assert result.suffix == ".docx"
        assert result.stat().st_size > 1000

    def test_docx_has_content(self, tmp_path):
        from core.doc_generator import generate_executive_docx
        from docx import Document

        out = tmp_path / "test.docx"
        generate_executive_docx(
            cards=_make_cards(),
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=3,
            output_path=out,
        )
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Cover elements
        assert "每日科技趨勢簡報" in full_text
        assert "Daily Tech Intelligence" in full_text

        # News card elements
        assert "Test News Title 1" in full_text
        assert "Test News Title 2" in full_text

        # QA block
        assert "Q" in full_text

        # Next steps
        assert "Next Steps" in full_text or "下一步" in full_text

    def test_docx_has_table(self, tmp_path):
        from core.doc_generator import generate_executive_docx
        from docx import Document

        out = tmp_path / "test.docx"
        generate_executive_docx(
            cards=_make_cards(),
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=3,
            output_path=out,
        )
        doc = Document(str(out))
        assert len(doc.tables) >= 1  # metrics table

    def test_empty_cards(self, tmp_path):
        from core.doc_generator import generate_executive_docx

        out = tmp_path / "empty.docx"
        result = generate_executive_docx(
            cards=[],
            health=_make_health(),
            report_time="2026-02-13 10:00",
            total_items=0,
            output_path=out,
        )
        assert result.exists()
        assert result.stat().st_size > 500


# ---------------------------------------------------------------------------
# Integration: generate_binary_reports
# ---------------------------------------------------------------------------


class TestBinaryReportsIntegration:
    def test_generate_binary_reports(self, tmp_path, monkeypatch):
        from core.education_renderer import generate_binary_reports

        # Monkey-patch project root so outputs go to tmp_path
        pptx_path, docx_path = generate_binary_reports(
            results=None,
            metrics=SAMPLE_METRICS,
            project_root=tmp_path,
        )

        assert pptx_path.exists()
        assert docx_path.exists()
        assert pptx_path.suffix == ".pptx"
        assert docx_path.suffix == ".docx"

    def test_with_structured_input(self, tmp_path):
        from core.education_renderer import generate_binary_reports

        results = [
            MergedResult(
                item_id="news_001",
                schema_a=SchemaA(
                    item_id="news_001",
                    title_zh="Test Title",
                    summary_zh="Test summary for the news item.",
                    category="tech",
                    entities=["CompanyA"],
                    key_points=["Key point 1"],
                    source_id="https://example.com",
                ),
                schema_b=SchemaB(item_id="news_001", final_score=8.0),
                schema_c=SchemaC(item_id="news_001"),
                passed_gate=True,
            ),
        ]
        report = DeepAnalysisReport(
            generated_at="2026-02-13",
            total_items=1,
            per_item_analysis=[
                ItemDeepDive(
                    item_id="news_001",
                    core_facts=["Fact 1"],
                    evidence_excerpts=["Evidence 1"],
                    signal_strength=0.8,
                ),
            ],
        )

        pptx_path, docx_path = generate_binary_reports(
            results=results,
            report=report,
            metrics=SAMPLE_METRICS,
            project_root=tmp_path,
        )

        assert pptx_path.exists()
        assert docx_path.exists()
        assert pptx_path.stat().st_size > 1000
        assert docx_path.stat().st_size > 1000
