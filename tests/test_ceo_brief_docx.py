"""CEO Brief DOCX generation tests.

Validates:
- DOCX generates successfully and can be opened
- Contains structured summary 5 section titles
- Contains Q1/Q2/Q3 text
- Contains data_card metric text
- Contains video_source and sources text
- Banned word / empty-talk guard
"""

from __future__ import annotations

import sys
from dataclasses import dataclass, field
from pathlib import Path
from unittest.mock import patch

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core.doc_generator import generate_executive_docx
from schemas.education_models import EduNewsCard, SystemHealthReport

# ---------------------------------------------------------------------------
# Banned words (same as existing tests)
# ---------------------------------------------------------------------------
BANNED_WORDS = [
    "ai捕捉", "AI Intel", "Z1", "Z2", "Z3", "Z4", "Z5",
    "pipeline", "ETL", "verify_run", "ingestion", "ai_core",
]


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _make_health() -> SystemHealthReport:
    return SystemHealthReport(
        success_rate=80.0,
        p50_latency=4.0,
        p95_latency=11.0,
        entity_noise_removed=2,
        total_runtime=30.0,
        run_id="test_docx_run",
        fail_reasons={"blocked": 1},
    )


def _make_event_card(**overrides) -> EduNewsCard:
    defaults = dict(
        item_id="test-001",
        is_valid_news=True,
        title_plain="Nvidia launches new GPU chip today",
        what_happened="Nvidia 今日推出新款 GPU 晶片，效能提升 40%，售價 2000 美元",
        why_important="影響 AI 訓練成本與效能",
        focus_action="評估是否採購",
        metaphor="就像汽車引擎升級，所有車都能跑更快",
        fact_check_confirmed=["Nvidia announced the chip at CES 2026",
                              "Performance benchmarks show 40% improvement"],
        evidence_lines=["Nvidia CEO Jensen Huang presented the new architecture"],
        technical_interpretation="新架構採用 4nm 製程",
        derivable_effects=["AI 訓練成本可能下降 20-30%", "競爭對手 AMD 壓力增加"],
        speculative_effects=["可能引發新一輪 AI 軍備競賽"],
        action_items=["評估新 GPU 對現有 AI 專案的影響"],
        video_suggestions=["Nvidia GPU chip launch 2026"],
        source_url="https://example.com/nvidia-gpu",
        category="人工智慧",
        final_score=8.5,
    )
    defaults.update(overrides)
    return EduNewsCard(**defaults)


def _extract_all_text(docx_path: Path) -> str:
    """Extract all text from a DOCX file."""
    from docx import Document
    doc = Document(str(docx_path))
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)
    return "\n".join(texts)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestDocxGeneration:
    def test_generates_and_opens(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "test.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            result = generate_executive_docx(
                cards, _make_health(), "2026-02-15 09:00", 5, out)
        assert result.exists()
        assert result.suffix == ".docx"
        assert result.stat().st_size > 0
        # Verify it can be opened
        from docx import Document
        doc = Document(str(result))
        assert len(doc.paragraphs) > 0

    def test_generates_with_empty_cards(self, tmp_path: Path):
        out = tmp_path / "empty.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            result = generate_executive_docx(
                [], _make_health(), "2026-02-15 09:00", 0, out)
        assert result.exists()


class TestStructuredSummary:
    def test_has_5_section_titles(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "summary.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        for section in ["AI Trends", "Tech Landing", "Market Competition",
                        "Opportunities & Risks", "Recommended Actions"]:
            assert section in text, f"Section '{section}' missing in DOCX"


class TestQAContent:
    def test_has_q1_q2_q3(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "qa.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        assert "Q1" in text, "Q1 missing"
        assert "Q2" in text, "Q2 missing"
        assert "Q3" in text, "Q3 missing"
        assert "WHY IT MATTERS" in text


class TestDataCard:
    def test_has_metric_text(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "dc.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        assert "Data Card" in text, "Data Card section missing"
        # Must contain numeric data
        assert any(c.isdigit() for c in text), "No numeric data found"


class TestVideoAndSources:
    def test_has_video_reference(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "vid.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        assert "Video" in text or "video" in text.lower(), "Video reference missing"

    def test_has_sources(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "src.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        assert "Source" in text or "source" in text.lower(), "Sources missing"


class TestBannedWords:
    def test_no_banned_words(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "bw.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        for bw in BANNED_WORDS:
            assert bw not in text, f"Banned word '{bw}' found in DOCX"


class TestChartSpec:
    def test_has_chart_spec(self, tmp_path: Path):
        cards = [_make_event_card()]
        out = tmp_path / "chart.docx"
        with patch("core.doc_generator.get_news_image", return_value=None):
            generate_executive_docx(cards, _make_health(), "2026-02-15 09:00", 5, out)
        text = _extract_all_text(out)
        assert "Chart Spec" in text or "Chart Type" in text, "Chart spec missing"
